import sys
from io import BytesIO
from pathlib import Path
import tempfile
import unittest

from docx import Document

APP_DIR = Path(__file__).resolve().parents[1] / "app"
if str(APP_DIR) not in sys.path:
    sys.path.insert(0, str(APP_DIR))

from modules.xml_word_engine_adapter import XMLWordEngineAdapter


class XMLWordEngineAdapterTests(unittest.TestCase):
    def _create_temp_doc(self):
        tmp_dir = tempfile.TemporaryDirectory()
        doc_path = Path(tmp_dir.name) / "plantilla.docx"
        return tmp_dir, doc_path

    def test_replace_variables_handles_split_marker(self):
        tmp_dir, doc_path = self._create_temp_doc()
        try:
            doc = Document()
            para = doc.add_paragraph()
            para.add_run("La relevancia es ")
            para.add_run("<<Editar en el word ")
            para.add_run("segun corresponda>>")
            para.add_run(" dentro del informe.")
            doc.save(doc_path)

            engine = XMLWordEngineAdapter(doc_path)
            engine.replace_variables({
                "<<Editar en el word segun corresponda>>": "Texto personalizado"
            })
            engine.clean_unused_markers()
            result_bytes = engine.get_document_bytes()
            engine.__del__()

            result_doc = Document(BytesIO(result_bytes))
            full_text = "\n".join(p.text for p in result_doc.paragraphs)

            self.assertIn("Texto personalizado", full_text)
            self.assertNotIn("<<Editar en el word segun corresponda>>", full_text)
        finally:
            tmp_dir.cleanup()

    def test_process_table_of_contents_removes_missing_entries(self):
        tmp_dir, doc_path = self._create_temp_doc()
        try:
            doc = Document()
            doc.add_paragraph("Índice <<Indice>>")
            doc.add_paragraph("Capítulo Uno <<1>>")
            doc.add_paragraph("Capítulo Fantasma <<2>>")
            doc.add_paragraph("<<fin Indice>>")
            doc.add_paragraph("<<1>> Capítulo Uno")
            doc.add_paragraph("Contenido capítulo uno")
            doc.save(doc_path)

            engine = XMLWordEngineAdapter(doc_path)
            engine.process_table_of_contents()
            engine.clean_unused_markers()
            result_bytes = engine.get_document_bytes()
            engine.__del__()

            result_doc = Document(BytesIO(result_bytes))
            paragraphs = [p.text for p in result_doc.paragraphs if p.text.strip()]
            joined = " ".join(paragraphs)

            # La entrada válida debe mantenerse con número de página
            self.assertTrue(any(
                "Capítulo Uno" in text and text.strip()[-1].isdigit()
                for text in paragraphs
            ))
            # La entrada inexistente debe eliminarse
            self.assertNotIn("Capítulo Fantasma", joined)
            # Los marcadores del índice deben eliminarse
            self.assertNotIn("<<Indice>>", joined)
            self.assertNotIn("<<fin Indice>>", joined)
        finally:
            tmp_dir.cleanup()


if __name__ == "__main__":
    unittest.main()
